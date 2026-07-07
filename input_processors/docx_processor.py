# ============================================================
# input_processors/docx_processor.py
# ============================================================
# STEP 5C: Microsoft Word (.docx) File Input Processor
#
# The user uploads a Word document containing meeting minutes,
# an agenda, or typed notes. This module uses 'python-docx'
# to iterate through all paragraphs and tables and extract
# the plain text.
#
# RESPONSIBILITY:
#   Read all paragraph text and table cell text from a .docx
#   file and return a clean, combined string for the LLM.
# ============================================================

import io
from docx import Document


def process_docx(file_bytes: bytes) -> str:
    """
    STEP 5C.1: Extract all text from a Word .docx file.

    Args:
        file_bytes (bytes): Raw bytes of the uploaded .docx file
                            (from Streamlit's st.file_uploader).

    Returns:
        str: Concatenated plain text from paragraphs and tables.
        Returns empty string if document has no text content.
    """
    # STEP 5C.1.1: Wrap bytes in BytesIO stream for python-docx
    docx_stream = io.BytesIO(file_bytes)

    # STEP 5C.1.2: Open the Word document
    try:
        doc = Document(docx_stream)
    except Exception as e:
        print(f"[docx_processor] Failed to open DOCX: {e}")
        return ""

    content_parts = []

    # STEP 5C.1.3: Extract text from all paragraphs
    # A paragraph in Word = a block of text separated by Enter key.
    # Headings, bullet points, and body text are all paragraphs.
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # Skip empty paragraphs (blank lines in Word)
            content_parts.append(text)

    # STEP 5C.1.4: Extract text from all tables
    # Word tables contain rows and cells; each cell may have text.
    # We flatten table content row by row, separating cells with " | ".
    for table in doc.tables:
        for row in table.rows:
            # Collect text from each cell in this row
            row_cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]
            if row_cells:
                # Join cells with a pipe separator (readable by LLM)
                content_parts.append(" | ".join(row_cells))

    # STEP 5C.1.5: Guard — warn if document was empty
    if not content_parts:
        print("[docx_processor] Warning: No text found in DOCX file.")
        return ""

    # STEP 5C.1.6: Join all parts with newlines
    return "\n".join(content_parts)
